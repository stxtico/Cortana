"""write_docx (PROMPTS.md A24) - "here's the file," not "here's some text you
can copy." Write-capable, gated the same way tools/write_file.py already is:
REQUIRES_CONFIRMATION = True, whitelisted to [tools].write_whitelist_dirs via
tools/_fs.py - the existing write whitelist, not a second path check invented
for this tool.
"""

from docx import Document

from tools import _fs

REQUIRES_CONFIRMATION = True

_WRITE_KEY = "write_whitelist_dirs"


def spec() -> dict:
    return {
        "type": "function",
        "function": {
            "name": "write_docx",
            "description": (
                "Create a Word document (.docx) with an optional title and body text "
                "(paragraphs separated by newlines). Only these directories (and their "
                f"subdirectories) are write-authorized: {_fs.whitelist_description(_WRITE_KEY)}. "
                "Requires user confirmation before it actually runs."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the .docx file, absolute or relative to the project root."},
                    "content": {"type": "string", "description": "Body text. Separate paragraphs with newlines."},
                    "title": {"type": "string", "description": "Optional heading at the top of the document."},
                },
                "required": ["path", "content"],
            },
        },
    }


def describe(path: str, content: str, title: str | None = None) -> str:
    preview = content if len(content) <= 120 else content[:120] + "…"
    heading = f" (titled {title!r})" if title else ""
    return f"Write a Word document{heading} to {path!r}:\n    {preview!r}"


async def execute(path: str, content: str, title: str | None = None) -> str:
    resolved = _fs.resolve_in_whitelist(path, _WRITE_KEY)
    resolved.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(resolved))

    return f"Wrote a Word document to {resolved}."
